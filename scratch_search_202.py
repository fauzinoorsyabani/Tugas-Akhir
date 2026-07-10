from docx import Document
import sys

def search_docx():
    path = r'd:\Code\Tugas Akhir\Skripsi\Laporan\Laporan Ujian Seminar Hasil\Draft Laporan Ujian Seminar Hasil_Fauzi Noorsyabani_227007042_(2).docx'
    try:
        doc = Document(path)
        print('=== HASIL PENCARIAN ANGKA 202 ===')
        found = False
        for i, p in enumerate(doc.paragraphs):
            if '202' in p.text:
                print(f'\n[Paragraf {i+1}]')
                print(f'Isi Teks: {p.text.strip()}')
                found = True
                
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if '202' in cell.text:
                        print(f'\n[Tabel ke-{t_idx+1}, Baris {r_idx+1}]')
                        print(f'Isi Sel: {cell.text.strip()}')
                        found = True

        if not found:
             print('TIDAK DITEMUKAN angka 202 di dalam dokumen.')
             
    except Exception as e:
        print(f'Error: {e}')

if __name__ == '__main__':
    search_docx()
